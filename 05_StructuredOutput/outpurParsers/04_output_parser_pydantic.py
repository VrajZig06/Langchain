from langchain_huggingface import HuggingFaceEndpoint,ChatHuggingFace
from langchain.output_parsers import PydanticOutputParser
from langchain.prompts import PromptTemplate
from pydantic import BaseModel,Field
from docx import Document


llm = HuggingFaceEndpoint(
 repo_id = "Qwen/Qwen3-Coder-480B-A35B-Instruct",
  task="text-generation",
  max_new_tokens=50,
  temperature=1.8
)

model = ChatHuggingFace(llm=llm)

class Person(BaseModel):
  name : str
  age : int = Field(gt=18)
  city : str

pasrer = PydanticOutputParser(pydantic_object=Person)

templet = PromptTemplate(
  template="""Give me name,age and cityname of frictional person in {place}.\n {formate_instruction}""",
  input_variables=['place'],
  partial_variables={"formate_instruction" : pasrer.get_format_instructions()} 
)

Country = input("Enter any Country Name here: ")

prompt = templet.invoke({
  "place": Country,
})

result = model.invoke(prompt)

final_result = pasrer.parse(result.content).model_dump()

doc = Document()
doc.add_heading("AI-Generated Report", level=1)
doc.add_paragraph(final_result)

# Save to file
doc.save("report.docx")