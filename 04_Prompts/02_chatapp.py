from langchain_huggingface import ChatHuggingFace,HuggingFaceEndpoint
import streamlit as st
from docx import Document

llm = HuggingFaceEndpoint(
  repo_id = "Qwen/Qwen3-Coder-480B-A35B-Instruct",
  task = "text-generation",
  max_new_tokens = 10
)

model = ChatHuggingFace(llm=llm)

# webiste 
st.header("Chat Bot using HuggingFace")
doc = Document()
doc.add_heading("Chat History")

if "User_Query" not in st.session_state:
    st.session_state.User_Query = []
if "AI_Response" not in st.session_state:
    st.session_state.AI_Response = []
if "Chat_Hostory" not in st.session_state:
    st.session_state.Chat_Hostory = []

for i in range(len(st.session_state.User_Query)):
    st.markdown(
        f"<div style='text-align: right;'><b style=' color: yellow;'>User:</b> {st.session_state.User_Query[i]}</div>",
        unsafe_allow_html=True
    )
    doc.add_paragraph(f"User : {st.session_state.User_Query[i]}")

    st.markdown(
        f"<div style='text-align: left; '><b style='color: green;'>AI:</b> {st.session_state.AI_Response[i]}</div>",
        unsafe_allow_html=True
    )
    doc.add_paragraph(f"AI : {st.session_state.AI_Response[i]}")

user_input = st.text_input("", key="user_input")
btn = st.button("Send")

if user_input.lower() == "exit":
    doc.save("chatHistory.docx")
    st.stop()

if btn and user_input:
    st.session_state.User_Query.append(user_input)
    st.session_state.Chat_Hostory.append(user_input)
    result = model.invoke(st.session_state.Chat_Hostory)
    st.session_state.AI_Response.append(result.content)
    st.session_state.Chat_Hostory.append(result.content)
    user_input = ""
    st.rerun()

